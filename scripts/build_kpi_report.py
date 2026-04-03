from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.79)
section.right_margin = Cm(2.79)

def h(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 1:
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    return p

def b(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = bold
    r.italic = italic
    if indent:
        p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p

def bul(doc, text, level=0):
    blts = ["  \u2022 ", "    \u25e6 ", "      \u25aa "]
    p = doc.add_paragraph()
    r = p.add_run(blts[level] + text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def alert(doc, level, title, items=None):
    clrs = {"CRITICAL": "C00000", "HIGH": "C05000", "MEDIUM": "BF8F00", "GOOD": "007000"}
    icons = {"CRITICAL": "\U0001F534", "HIGH": "\U0001F7E0", "MEDIUM": "\U0001F7E1", "GOOD": "\U0001F7E2"}
    p = doc.add_paragraph()
    r = p.add_run(icons.get(level, "") + " " + level + ": " + title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(clrs.get(level, "000000"))
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    if items:
        for item in items:
            bul(doc, item)

def finding(doc, num, title, desc):
    p = doc.add_paragraph()
    r = p.add_run("FINDING " + str(num) + ": ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    b(doc, desc)

def tbl(doc, headers, rows):
    ncols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, hdr in enumerate(headers):
        c = t.cell(0, i)
        clr_cell(c, hdr, bold=True, white=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            c = t.cell(r_idx + 1, c_idx)
            clr_cell(c, val, bold=False, white=False, alt=(r_idx % 2 == 1))
    doc.add_paragraph()
    return t

def clr_cell(cell, text, bold=False, white=False, alt=False):
    for par in cell.paragraphs:
        for rn in par.runs:
            rn.text = ""
        rn = (par.runs[0] if par.runs else par.add_run())
        rn.text = text
        rn.font.size = Pt(9.5)
        rn.bold = bold
        if white:
            rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd_el = cell._element.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    if white:
        shd_el.set(qn("w:fill"), "1F4E79")
    else:
        shd_el.set(qn("w:fill"), "F2F2F2" if alt else "FFFFFF")
    shd_el.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shd_el)

# ===== COVER PAGE =====
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHIPCHIP FINANCE KPI REPORT")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("COMPREHENSIVE REVIEW & OPTIMIZATION RECOMMENDATIONS")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Based on: Finance KPI Report \u2014 Week 13, 2026 (Mar 19\u201325)")
r.font.size = Pt(11)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared: April 1, 2026")
r.font.size = Pt(11)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL \u2014 FOR INTERNAL USE ONLY")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===== TOC =====
h(doc, "Table of Contents", 1)
toc_items = [
    "Section 1: Executive Summary",
    "Section 2: Complete Page-by-Page Analysis of Current Report",
    "Section 3: 7 Key Findings",
    "Section 4: Proposed Report Structure",
    "Section 5: Section-by-Section Replacement Guide",
    "Section 6: Implementation Roadmap (2-Week Sprint)",
    "Section 7: Appendix \u2014 Full Proposed Report Template",
    "Closing Recommendations",
]
for i, item in enumerate(toc_items):
    p = doc.add_paragraph()
    r = p.add_run(str(i + 1) + ".  " + item)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
doc.add_page_break()

# ===== SECTION 1 =====
h(doc, "Section 1: Executive Summary", 1)
b(doc, "The current Finance KPI Report (Week 13, 2026 \u2014 Mar 19\u201325) is 4 pages long but contains ZERO financial data. It is structured as a meeting discussion log rather than a performance dashboard. This document provides a complete, evidence-based review with a ready-to-implement replacement template.")
h(doc, "What We Found", 2)
finding(doc, 1, "No Financial Metrics", "The report contains no revenue figures, cost data, margin analysis, cash position, or any quantitative KPI. The only section with content is Grants, which tracks process status \u2014 not numbers.")
finding(doc, 2, "Meetings Tracked, Not Performance", "Every section follows Discussion Point \u2192 Responsible \u2192 Outcome \u2192 Timeline. That is meeting minutes format, not KPI format.")
finding(doc, 3, "Deadlines Missed and Undefined", "Multiple action items carry 'Undefined' as their timeline and past-due deadlines. The report cannot track accountability without this.")
h(doc, "What We Recommend", 2)
b(doc, "Replace the current structure with a single-page KPI dashboard covering 5 sections: Revenue & Margins, Cash & Costs, Unit Economics, Grants & Funding, Top Risks & Actions. Use traffic-light colours on numbers (not on discussions). Close the loop by showing last week's actions as Done / In Progress / Overdue.")
alert(doc, "HIGH", "Report is structurally broken \u2014 fix before next reporting cycle", [
    "Without numbers, the report cannot fulfil any of its 5 stated goals",
    "The current format consumes time but provides zero analytical value",
    "A 1-page replacement will be faster to prepare AND more useful",
])
doc.add_page_break()

# ===== SECTION 2 =====
h(doc, "Section 2: Complete Page-by-Page Analysis", 1)

h(doc, "Page 1 \u2014 Title & Goals", 2)
b(doc, "Contains the title, Deming quote, report period, and 5 goals. No data on this page.", bold=True)
b(doc, "Goals stated: Profitability, Cross-check efficiency, Reduce costs, Ensure data accuracy, Detect fraud")
alert(doc, "MEDIUM", "Goals page is unnecessary padding", [
    "A Deming quote does not replace data",
    "Goals should be implicit in the KPI design \u2014 not listed separately",
    "Recommendation: Remove entirely. Use a one-line title bar.",
])

h(doc, "Page 2 \u2014 Grants (Bealu)", 2)
b(doc, "Structure: 5 bullet points \u2014 Disbursement, Pipeline, Active Opportunities, KPI tracker, Next Actions. No numbers provided.", bold=True)
alert(doc, "HIGH", "Grants section is hollow \u2014 no numbers", [
    "'Grant Disbursement Amount' \u2014 no amount listed",
    "'Pipeline Status' \u2014 no pipeline value, no conversion rate",
    "'Active Opportunities' \u2014 no count, no expected value",
    "'KPI Tracker (Mesirat, CASA)' \u2014 no KPI values",
    "Without numbers, this is a task list, not a financial section",
])

h(doc, "Page 3 \u2014 Discussion Points (Part 1)", 2)
b(doc, "Table: Last weeks' discussion points \u2192 Funding", bold=True)
bul(doc, "Slide review: Team voted keep (green) vs remove (red) \u2014 deadline Apr 2, outcome pending")
bul(doc, "Data discrepancy: Bank system overload, timezone mismatch (ChipChip=UTC) \u2014 owners: Liliane/Bealu/Kira \u2014 timeline: Undefined")
alert(doc, "CRITICAL", "'Undefined' is the most unacceptable status in a KPI report", [
    "An action item with no deadline is a wish, not an action item",
    "Timeline = Undefined on a data accuracy issue means the root cause is NOT being addressed",
    "This is the highest-priority finding in the entire document",
])

h(doc, "Page 4 \u2014 Discussion Points (Part 2) & Actions", 2)
b(doc, "Second half of the discussion log:", bold=True)
bul(doc, "KPI voting sheet \u2014 traffic light system for slides \u2014 deadline Mar 26")
bul(doc, "Cash flow slide improvement \u2014 add variance, target, error/date \u2014 owner: Befeker \u2014 Mar 26")
bul(doc, "CP1 & CP2 \u2014 source data from payment system (Supply Chain) \u2014 NO owner listed")
bul(doc, "Logistics tracking \u2014 push data directly from app \u2014 NO owner listed")
alert(doc, "HIGH", "Action items lack owners and completion status", [
    "CP1/CP2 data source change: no owner assigned",
    "Logistics app tracking: no owner assigned",
    "No Done column \u2014 impossible to know if Mar 26 actions were completed",
])
doc.add_page_break()

# ===== SECTION 3 =====
h(doc, "Section 3: 7 Key Findings", 1)

fdata = [
    ("ZERO FINANCIAL NUMBERS IN A FINANCE REPORT",
     "The report title is 'Finance KPI Report' but contains no revenue, COGS, margin, cash position, unit economics, burn rate, or variance analysis. Grants is the only section and it tracks process items, not financial outcomes. KPI stands for Key Performance Indicator. Without indicators, the report has no KPIs.",
     "Foundational problem \u2014 makes every other issue secondary"),
    ("REPORT IS A MEETING LOG, NOT A DASHBOARD",
     "Every section is Discussion Point \u2192 Responsible \u2192 Outcome \u2192 Timeline. This is meeting minutes format. Example: 'CP1 margin is 18% vs 22% target' (KPI) vs. 'Team discussed whether to source CP1 from Supply Chain' (meeting log). The document consistently uses the latter.",
     "No decision-making value"),
    ("CRITICAL DATA DISCREPANCY WITH NO RESOLUTION",
     "Data discrepancy between ChipChip records and bank confirmations. Causes: (a) banks not sending payment confirmations (system overload), (b) timezone mismatch (ChipChip=UTC vs bank local TZ). Status: Undefined.",
     "Highest priority \u2014 data integrity failure"),
    ("SLIDES BEING DESIGNED FOR DESIGN SAKE",
     "Significant time spent on slide voting (green/red) \u2014 effort goes to aesthetics, not financial data. The question should be 'what numbers do we need?' not 'which slides?'",
     "Misplaced effort"),
    ("CP1 / CP2 DATA SOURCE UNCLEAR",
     "Page 4 says CP1 and CP2 should be taken from the payment system. No owner assigned, no deadline set, unclear if the system has this data. These are the most important KPIs for the business.",
     "Critical metrics orphaned"),
    ("ACTION ITEMS LACK OWNERS AND TRACKING",
     "CP1/CP2 data source: no owner. Logistics app tracking: no owner. Deadlines are past (Mar 26). No Done column. If you cannot tell which last week actions are complete, the report fails accountability.",
     "Broken accountability loop"),
    ("REPORT COSTS MORE TIME THAN IT SAVES",
     "4 pages of meeting notes that take time to compile but provide no decision-ready data. A 1-page KPI with traffic-light numbers takes 10 minutes to read and 20 minutes to prepare.",
     "Negative ROI on reporting time"),
]
for title, desc, tag in fdata:
    h(doc, "Finding: " + title)
    b(doc, desc)
    p = doc.add_paragraph()
    r = p.add_run("Impact: " + tag)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
    doc.add_paragraph()
doc.add_page_break()

# ===== SECTION 4 =====
h(doc, "Section 4: Proposed Report Structure", 1)
b(doc, "The replacement report should be exactly ONE PAGE (front). A second page for risk details is acceptable \u2014 the KPI summary must be scannable in 60 seconds.", bold=True)

h(doc, "Design Principles", 2)
for t in ["Every section must contain at least ONE number",
           "Traffic lights on metrics, not on discussions",
           "Show last week's actions: Done / In Progress / Overdue",
           "No quotes, no goals lists, no philosophy \u2014 just data",
           "One owner for the entire report (rotate monthly)",
           "Distribute 24 hours before weekly call \u2014 10 min read, 50 min discussion"]:
    bul(doc, t)

h(doc, "Proposed Sections", 2)
tbl(doc, ["Section", "What It Covers", "Frequency"], [
    ["1. Revenue & Margins", "Actual Birr figures with targets and variances", "Weekly"],
    ["2. Cash & Operating Costs", "Cash position, burn rate, top cost variances", "Weekly"],
    ["3. Unit Economics", "CP1, CP2, cost per delivery, cost per order", "Weekly"],
    ["4. Grants & Funding", "Disbursed, pipeline, active opportunities with amounts", "Weekly"],
    ["5. Top 3 Risks & Actions", "What is broken, who owns it, status", "Weekly"],
])

h(doc, "Visual Design", 2)
bul(doc, "One-page grid layout \u2014 dashboard card style")
bul(doc, "Each metric: Value | Target | Variance | Status")
bul(doc, "Data Confidence score: [92%] or [\u26a0\ufe0f 78% \u2014 see Note 1]")
bul(doc, "ChipChip brand colours \u2014 white/red/black, clean layout")
bul(doc, "PDF export only \u2014 prevent tampering")
doc.add_page_break()

# ===== SECTION 5 =====
h(doc, "Section 5: Section-by-Section Replacement Guide", 1)

h(doc, "REMOVE", 2)
for r in ["Page 1 (Deming quote, goals list) \u2014 DELETE",
           "KPI Voting Sheet table \u2014 move to meeting notes only",
           "All Discussion Points tables \u2014 belong in meeting minutes",
           "Slide Review voting results \u2014 irrelevant to performance",
           "Any section without a financial number"]:
    alert(doc, "HIGH", r)

h(doc, "TRANSFORM", 2)
for title, desc in [
    ("Grants section", "Add: Disbursement Birr, pipeline value, conversion rate. Remove: bullet-task lines without numbers."),
    ("CP1/CP2 discussion", "Stop debating data sources. Define formulas. Automate extraction. Report: CP1 = X% (target Y%)."),
    ("Data discrepancy issue", "Move to Risk #1 in Top Risks. Assign owner + deadline. Track weekly until resolved."),
    ("Cash flow slide suggestion", "Build it. Add line: Cash: X Birr | Variance: +/-Y% | Trend: Up/Down."),
]:
    b(doc, title + ":", bold=True)
    b(doc, desc, indent=True)

h(doc, "ADD (Currently Missing)", 2)
tbl(doc, ["Metric", "Missing Value", "Data Source"], [
    ["Total Revenue", "Birr figure, B2B/B2C split", "E-commerce platform / payment system"],
    ["COGS", "Cost by product line", "Inventory system / supplier invoices"],
    ["CP1 (Gross Margin)", "Revenue minus COGS as %", "Automate from revenue + COGS"],
    ["CP2 (Net Margin)", "Revenue minus COGS minus warehouse minus delivery", "Add warehouse/delivery to CP1 calc"],
    ["Cash Position", "Current bank / float balance", "Bank statements / aggregator"],
    ["Burn Rate", "Weekly cash consumption", "OpEx tracker"],
    ["Cost Variance", "Top 3 items over budget", "Expense ledger vs budget"],
    ["Grant Disbursement", "Actual ETB received this week", "Grant tracking (Bealu)"],
    ["Confidence Score", "% data accuracy", "Auto-calc from source status"],
])
doc.add_page_break()

# ===== SECTION 6 =====
h(doc, "Section 6: Implementation Roadmap", 1)
h(doc, "2-Week Sprint", 2)

h(doc, "WEEK 1 \u2014 Data Foundation", 2)
for d in ["Day 1-2: Define CP1, CP2, unit metric formulas (single owner)",
           "Day 3: Build data extraction from payment system / Supply Chain",
           "Day 4-5: Populate first version of 1-page template with real numbers",
           "Day 5: Review with finance team \u2014 validate accuracy"]:
    bul(doc, d)

h(doc, "WEEK 2 \u2014 Process & Automation", 2)
for d in ["Day 1-2: Build automated report generation (spreadsheet to PDF)",
           "Day 3: Traffic-light logic: On target / +/-10% / Beyond",
           "Day 4: Confidence indicator: flag unreliable sources",
           "Day 5: Train team on new format; kill old report",
           "Day 7: First production run"]:
    bul(doc, d)

alert(doc, "HIGH", "3 things to do THIS WEEK", [
    "Assign a single Report Owner (one person accountable for compilation)",
    "Define CP1 and CP2 formulas and their data sources \u2014 write them down",
    "Populate new 1-page template with at least Revenue, Cash, Grant numbers for Week 14",
])
doc.add_page_break()

# ===== SECTION 7: APPENDIX =====
h(doc, "Appendix: Complete Proposed Report Template", 1)
b(doc, "Use this exact template for every weekly Finance KPI report going forward.", bold=True)
doc.add_paragraph()

b(doc, "CHIPCHIP FINANCE KPI REPORT \u2014 Week [XX], 2026", bold=True)
b(doc, "[Date Range] | Prepared: [Name] | Distribution: Finance Team & Leadership")
b(doc, "Data Confidence: [95%] | Prev Week Actions: 2/3 Done | Overdue: 1")
doc.add_paragraph()

h(doc, "1. Revenue & Margins", 2)
tbl(doc, ["Metric", "This Week", "Target", "Variance", "Status"], [
    ["Total Revenue", "[___] Birr", "[___] Birr", "[___]", "[ ]"],
    ["B2B Revenue", "[___] Birr", "[___] Birr", "[___]", "[ ]"],
    ["B2C Revenue", "[___] Birr", "[___] Birr", "[___]", "[ ]"],
    ["CP1 (Gross Margin)", "[___]%", "[___]%", "[___] pp", "[ ]"],
    ["CP2 (Net Margin)", "[___]%", "[___]%", "[___] pp", "[ ]"],
])

h(doc, "2. Cash & Operating Costs", 2)
tbl(doc, ["Metric", "Value", "Notes", "Status"], [
    ["Cash Position", "[___] Birr", "As of End of Week", "[ ]"],
    ["Weekly Burn Rate", "[___] Birr", "OpEx only", "[ ]"],
    ["Top Cost Variance", "[___] Birr", "[Item name]", "[ ]"],
    ["Warehouse Costs", "[___] Birr", "vs budget", "[ ]"],
    ["Delivery Costs", "[___] Birr", "Avg X Birr / order", "[ ]"],
])

h(doc, "3. Unit Economics", 2)
tbl(doc, ["Metric", "Value", "Target", "Status"], [
    ["Avg Order Value", "[___] Birr", "[___] Birr", "[ ]"],
    ["Delivery Cost / Order", "[___] Birr", "3 Birr target", "[ ]"],
    ["Cost Per Acquisition", "[___] Birr", "[___] Birr", "[ ]"],
    ["Orders / Week", "[___]", "[___]", "[ ]"],
])

h(doc, "4. Grants & Funding (Owner: Bealu)", 2)
tbl(doc, ["Metric", "Value", "Notes", "Status"], [
    ["Disbursed This Week", "[___] Birr", "[Grant name]", "[ ]"],
    ["Pipeline (Committed)", "[___] Birr", "[X] opportunities", "[ ]"],
    ["Active Opportunities", "[#]", "Expected: [___] Birr", "[ ]"],
    ["KPI Tracker Score", "[___]", "Mesirat / CASA", "[ ]"],
    ["Next Key Action", "[1 line]", "Due: [date]", "[ ]"],
])

h(doc, "5. Top 3 Risks & Outstanding Actions", 2)
tbl(doc, ["#", "Risk / Action", "Owner", "Deadline", "Status"], [
    ["1", "[Description]", "[Name]", "[Date]", "[ ]"],
    ["2", "[Description]", "[Name]", "[Date]", "[ ]"],
    ["3", "[Description]", "[Name]", "[Date]", "[ ]"],
])
doc.add_paragraph()

h(doc, "Notes", 2)
for n in [
    "Status: On Track / Warning (+/-10%) / Critical (beyond +/-10% or overdue)",
    "Data Confidence: Percentage of data sources with clean data. Below 85% = flag",
    "Previous Week Actions: Summarise last week's items before adding new ones",
    "Maximum 1 page per report. Detail goes to separate appendix",
    "Distribute every Monday by 10:00 to all call participants",
]:
    bul(doc, n)
doc.add_page_break()

# ===== CLOSING =====
h(doc, "Final Recommendations Summary", 1)
b(doc, "The current report is well-intentioned but fundamentally misaligned with its purpose: it tracks discussions instead of metrics, uses narrative format instead of data presentation, and has no financial numbers in a document titled 'Finance KPI Report.'")
doc.add_paragraph()
b(doc, "The fix is structural, not incremental: replace the format, not the template. The proposed 1-page dashboard contains everything needed for a productive 1-hour weekly KPI call while eliminating all waste.")
doc.add_paragraph()

h(doc, "3-Step Action Plan", 2)
b(doc, "Step 1 (This Week): Define CP1/CP2 formulas and data sources. Assign a single Report Owner.")
b(doc, "Step 2 (Next Week): Build and populate the 1-page template with real numbers for Week 14.")
b(doc, "Step 3 (Week 3): Automate the data pipeline. Run first production report. Retire the old format entirely.")
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('"In God we trust. All others must bring data." \u2014 W. Edwards Deming')
r.italic = True
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

out = "/home/ubuntu/.openclaw/workspace/deliverables/Finance_KPI_Report_Review_and_Optimization_Recommendations.docx"
doc.save(out)
print("SAVED: " + out)
